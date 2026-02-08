import io
import base64
from datetime import date, timedelta
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


# =============================
# CONFIG
# =============================

st.set_page_config(page_title="MDC - Quotation System", layout="wide")

WORD_TEMPLATE_PATH = r"templates/PQ Template.docx"
LOGO_PATH = r"assets/logo.png"

COMPANY_NAME = "Mosaic Design Corporation (MDC)"
SIDEBAR_SUBTITLE = "Quotation System"


# =============================
# GLOBAL STYLES
# =============================

st.markdown(
    """
    <style>
      html, body, [class*="css"] { direction: ltr; }
      section.main, section.main * { direction: rtl; }
      section[data-testid="stSidebar"], section[data-testid="stSidebar"] * { direction: ltr; }

      .block-container { padding-top: 1.2rem; }
      h1, h2, h3, p, label { text-align: right; }

      .stTextInput input, .stNumberInput input, .stDateInput input, .stTextArea textarea {
        border-radius: 14px !important;
        padding: 10px 12px !important;
      }

      .card {
        background: rgba(255,255,255,0.70);
        border: 1px solid rgba(0,0,0,0.06);
        border-radius: 18px;
        padding: 16px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.04);
        margin-bottom: 14px;
      }

      .section-title {
        font-size: 22px;
        font-weight: 800;
        margin-bottom: 10px;
      }

      .brand-header {
        display: flex;
        align-items: center;
        gap: 12px;
        background: #ffffff;
        border-radius: 18px;
        padding: 14px;
        box-shadow: 0 10px 26px rgba(0,0,0,0.08);
        border: 1px solid rgba(0,0,0,0.06);
        margin-bottom: 14px;
      }

      .brand-icon {
        width: 46px;
        height: 46px;
        border-radius: 14px;
        background: #f4f6f8;
        display: flex;
        align-items: center;
        justify-content: center;
        overflow: hidden;
        border: 1px solid rgba(0,0,0,0.08);
        flex-shrink: 0;
      }

      .brand-title {
        color: #0b2c4d;
        font-size: 20px;
        font-weight: 900;
      }

      .brand-sub {
        color: #334e68;
        font-size: 13px;
        font-weight: 600;
      }
    </style>
    """,
    unsafe_allow_html=True
)


# =============================
# UTIL: Logo
# =============================

def logo_html(path: str):
    try:
        b = Path(path).read_bytes()
        b64 = base64.b64encode(b).decode()
        return f'<img src="data:image/png;base64,{b64}" style="width:32px;height:32px;object-fit:contain;" />'
    except Exception:
        return "MDC"


# =============================
# SIDEBAR
# =============================

with st.sidebar:
    st.markdown(
        f"""
        <div class="brand-header">
            <div class="brand-icon">{logo_html(LOGO_PATH)}</div>
            <div>
                <div class="brand-title">{COMPANY_NAME}</div>
                <div class="brand-sub">{SIDEBAR_SUBTITLE}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    page = st.radio(
        "Navigation",
        ["Price Quotation", "Settings (Soon)", "Reports (Soon)"],
        index=0
    )


# =============================
# WORD HELPERS
# =============================

def replace_in_paragraph(paragraph, mapping):
    for run in paragraph.runs:
        for k, v in mapping.items():
            token = f"{{{{{k}}}}}"
            if token in run.text:
                run.text = run.text.replace(token, str(v))

    # fallback لو placeholder متقسم
    full = "".join(r.text for r in paragraph.runs)
    new = full
    for k, v in mapping.items():
        new = new.replace(f"{{{{{k}}}}}", str(v))

    if new != full:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""


def replace_in_doc(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)


def find_items_table(doc):
    keys = ["التوصيف", "الكمية", "سعر الوحدة", "سعر البند"]
    for t in doc.tables:
        header = " ".join(c.text for c in t.rows[0].cells)
        if any(k in header for k in keys):
            return t
    return None


def fill_items_table(doc, df):
    table = find_items_table(doc)
    if not table:
        return False

    headers = [c.text for c in table.rows[0].cells]

    def idx(k):
        for i, h in enumerate(headers):
            if k in h:
                return i
        return None

    i_no = idx("م")
    i_desc = idx("التوصيف")
    i_qty = idx("الكمية")
    i_unit = idx("سعر الوحدة")
    i_total = idx("سعر البند")
    i_notes = idx("ملاحظات")

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for i, row in df.iterrows():
        cells = table.add_row().cells
        if i_no is not None: cells[i_no].text = str(row["م"])
        if i_desc is not None: cells[i_desc].text = str(row["التوصيف"])
        if i_qty is not None: cells[i_qty].text = str(row["الكمية"])
        if i_unit is not None: cells[i_unit].text = str(row["سعر الوحدة"])
        if i_total is not None: cells[i_total].text = str(row["سعر البند"])
        if i_notes is not None: cells[i_notes].text = str(row["ملاحظات"])

    return True


# =============================
# PAGE: Price Quotation
# =============================

def render_price_quotation():
    st.markdown(
        f"""
        <div class="brand-header">
            <div class="brand-icon">{logo_html(LOGO_PATH)}</div>
            <div>
                <div class="brand-title">{COMPANY_NAME}</div>
                <div class="brand-sub">Price Quotation</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    with st.container():
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="section-title">🧷 بيانات عرض السعر</div>', unsafe_allow_html=True)

        po_no = st.text_input("رقم طلب الشراء")
        quotation_no = st.text_input("عرض سعر رقم")
        client_name = st.text_input("السادة شركة")
        attn_engineer = st.text_input("عناية المهندس")
        quotation_date = st.date_input("تاريخ عرض السعر", value=date.today())

        c1, c2 = st.columns(2)
        with c1:
            delivery_days = st.number_input("مدة التوريد (يوم)", 1, 365, 30)
        with c2:
            validity_days = st.number_input("مدة الارتباط بالسعر (يوم)", 1, 365, 20)

        expiry_date = quotation_date + timedelta(days=int(validity_days))

        st.text_input(
            "تاريخ انتهاء عرض السعر (محسوب تلقائيًا)",
            value=expiry_date.strftime("%Y/%m/%d"),
            disabled=True
        )

        delivery_text = f"مدة التوريد {delivery_days} يوم من تاريخ استلام العينات من الشركة"
        validity_text = f"مدة الارتباط بالاسعار {validity_days} يوم حتى تاريخ {expiry_date.strftime('%Y/%m/%d')}"

        st.markdown('</div>', unsafe_allow_html=True)

    # Items
    if "items" not in st.session_state:
        st.session_state["items"] = pd.DataFrame(
            [{"م": 1, "التوصيف": "", "الكمية": 1, "سعر الوحدة": 0.0, "ملاحظات": ""}]
        )

    with st.form("items_form"):
        df = st.data_editor(
            st.session_state["items"],
            num_rows="dynamic",
            use_container_width=True
        )
        if st.form_submit_button("💾 حفظ البنود"):
            st.session_state["items"] = df

    calc = st.session_state["items"].copy()
    calc["سعر البند"] = calc["الكمية"] * calc["سعر الوحدة"]
    subtotal = calc["سعر البند"].sum()

    st.metric("إجمالي البنود", f"{subtotal:,.2f}")

    notes = st.text_area("ملاحظات إضافية", height=120)

    if st.button("Generate Word"):
        doc = Document(WORD_TEMPLATE_PATH)

        mapping = {
            "PO_NO": po_no,
            "QUOTATION_NO": quotation_no,
            "CLIENT_NAME": client_name,
            "ATTN_ENGINEER": attn_engineer,
            "QUOTATION_DATE": quotation_date.strftime("%Y/%m/%d"),
            "EXPIRY_DATE": expiry_date.strftime("%Y/%m/%d"),
            "DELIVERY_TEXT": delivery_text,
            "VALIDITY_TEXT": validity_text,
            "SUBTOTAL": f"{subtotal:,.2f}",
            "NOTES_BOX": notes,
        }

        replace_in_doc(doc, mapping)
        fill_items_table(doc, calc)

        buf = io.BytesIO()
        doc.save(buf)
        st.download_button(
            "⬇️ تحميل Word",
            buf.getvalue(),
            file_name=f"Quotation_{quotation_no or 'NA'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


# =============================
# ROUTER
# =============================

if page == "Price Quotation":
    render_price_quotation()
else:
    st.info("الصفحة دي لسه تحت التطوير 👷‍♂️")
